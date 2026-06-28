import io
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def build_translated_docx(translated_text: str, output_path: str, target_language: str = "ar") -> str:
    """Render plain translated text into a clean .docx (fallback path)."""
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)

    is_rtl = target_language.lower() in {"ar", "arabic", "he", "hebrew", "fa", "farsi", "ur", "urdu"}

    for raw_line in translated_text.split("\n"):
        line = raw_line.strip()
        paragraph = document.add_paragraph(line)
        if is_rtl:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.right_to_left = True

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def parse_layout_plan(raw: str) -> dict:
    """Parse the model's layout_plan_json, tolerating ```json fences or stray
    text around the JSON object. Raises ValueError if no JSON object is found.
    """
    if not raw:
        raise ValueError("empty layout_plan response")
    text = raw.strip()
    # Strip a leading/trailing markdown code fence if present.
    fence = re.match(r"^```(?:json)?\s*(.*?)\s*```$", text, re.DOTALL)
    if fence:
        text = fence.group(1).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    # Last resort: grab the outermost { ... } span and try again.
    start, end = text.find("{"), text.rfind("}")
    if start != -1 and end != -1 and end > start:
        return json.loads(text[start : end + 1])
    raise ValueError("no valid JSON object in layout_plan response")


def _apply_font(run, name: str, size_pt: int, *, bold: bool = False) -> None:
    """Apply font to a run including the complex-script (Arabic) attributes,
    which python-docx's run.font.name/size do NOT set on their own. Without the
    cs font/size, Arabic text ignores the requested font (e.g. Sakkal Majalla).
    """
    run.bold = bold
    run.font.name = name
    run.font.size = Pt(size_pt)
    rpr = run._r.get_or_add_rPr()

    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), name)

    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(int(size_pt) * 2))  # half-points
    rpr.append(sz_cs)

    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rpr.append(rtl)


def _rtl_paragraph(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.insert(0, bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_text_block(doc, text: str, font: str, size: int, *, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text or "")
    _apply_font(run, font, size, bold=bold)
    _rtl_paragraph(paragraph)


def build_docx_from_layout_plan(layout_plan: dict) -> bytes:
    """Build a professional RTL Arabic DOCX from a layout_plan_json dict and
    return the file bytes. Honors font_family/font_size from the plan and sets
    proper complex-script attributes so Arabic renders in the chosen font."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    font_family = layout_plan.get("font_family") or "Sakkal Majalla"
    try:
        font_size = int(str(layout_plan.get("font_size", "14")).strip() or "14")
    except (TypeError, ValueError):
        font_size = 14

    style = doc.styles["Normal"]
    style.font.name = font_family
    style.font.size = Pt(font_size)

    size_for = {
        "title": font_size + 4,
        "subtitle": font_size + 2,
        "section_heading": font_size + 1,
    }

    for block in layout_plan.get("blocks", []) or []:
        btype = block.get("type", "paragraph")
        content = block.get("content", "")

        if btype in ("title", "subtitle", "section_heading"):
            _add_text_block(doc, content, font_family, size_for[btype], bold=True)

        elif btype in ("paragraph", "signature_block"):
            if btype == "signature_block":
                doc.add_paragraph()  # spacing before the signature
            _add_text_block(doc, content, font_family, font_size)

        elif btype == "field_table":
            rows = block.get("rows", []) or []
            if rows:
                table = doc.add_table(rows=len(rows), cols=2)
                table.style = "Table Grid"
                for i, row_data in enumerate(rows):
                    # RTL: value on the right (col 0), field label on the left (col 1).
                    value_cell = table.rows[i].cells[0]
                    field_cell = table.rows[i].cells[1]
                    for cell, value, bold in (
                        (value_cell, row_data.get("value", ""), False),
                        (field_cell, row_data.get("field", ""), True),
                    ):
                        para = cell.paragraphs[0]
                        run = para.add_run(str(value))
                        _apply_font(run, font_family, font_size, bold=bold)
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif btype == "data_table":
            headers = block.get("headers", []) or []
            rows = block.get("rows", []) or []
            if headers:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = "Table Grid"
                for j, header in enumerate(reversed(headers)):  # reversed for RTL
                    para = table.rows[0].cells[j].paragraphs[0]
                    run = para.add_run(str(header))
                    _apply_font(run, font_family, font_size, bold=True)
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for i, row_data in enumerate(rows):
                    for j, val in enumerate(reversed(list(row_data))):
                        para = table.rows[i + 1].cells[j].paragraphs[0]
                        run = para.add_run(str(val))
                        _apply_font(run, font_family, font_size)
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif btype == "page_break":
            doc.add_page_break()

        elif btype == "spacer":
            doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
