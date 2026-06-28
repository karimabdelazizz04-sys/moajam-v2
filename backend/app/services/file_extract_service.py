import gc
from io import BytesIO
from pathlib import Path

import docx
from pypdf import PdfReader

# Below this many characters, a PDF is treated as scanned/image-only and routed
# through Claude Vision OCR instead of trusting the (empty) embedded text layer.
_MIN_PDF_TEXT_CHARS = 20

# OCR runs on a 512MB free tier, so keep memory tight: render at a modest DPI,
# one page at a time, and cap how many pages we'll process per document.
_OCR_DPI = 120
_OCR_MAX_PAGES = 10

# Vision-based translation: render pages a bit sharper than OCR (Claude reads
# the whole page), as JPEG to stay under Anthropic's per-image/request size
# limits, capped so the request stays within the token/size budget.
_VISION_DPI = 150
_VISION_MAX_PAGES = 10


class UnsupportedFileType(Exception):
    pass


def extract_text(path: str, ocr_fallback: bool = True) -> str:
    """Extract raw text from an uploaded .docx, .pdf or .txt file.

    For PDFs, `ocr_fallback=False` returns whatever the embedded text layer
    yields (even if empty) without the expensive Claude-Vision OCR pass - used
    when the text is only needed for cheap routing/retrieval, not as content.
    """
    suffix = Path(path).suffix.lower()

    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".pdf":
        return _extract_pdf(path, ocr_fallback=ocr_fallback)
    if suffix == ".txt":
        return Path(path).read_text(encoding="utf-8", errors="ignore")

    raise UnsupportedFileType(f"Unsupported file type: {suffix}")


def _extract_docx(path: str) -> str:
    document = docx.Document(path)
    parts: list[str] = []

    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_pdf(path: str, ocr_fallback: bool = True) -> str:
    reader = PdfReader(path)
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    if len(text.strip()) >= _MIN_PDF_TEXT_CHARS or not ocr_fallback:
        return text
    # Empty/near-empty text layer -> the PDF is almost certainly scanned. Fall
    # back to rendering each page to an image and OCR'ing it with Claude Vision.
    return _ocr_pdf(path)


def render_pdf_to_images(
    path: str,
    dpi: int = _VISION_DPI,
    max_pages: int = _VISION_MAX_PAGES,
    quality: int = 85,
) -> tuple[list[bytes], int]:
    """Render up to `max_pages` PDF pages to JPEG bytes for a Claude Vision
    request, one page at a time to keep memory low.

    Returns (images, total_pages) where total_pages is the real page count so
    the caller can note when a long document was truncated.
    """
    from pdf2image import convert_from_path, pdfinfo_from_path

    try:
        total_pages = pdfinfo_from_path(path)["Pages"]
    except Exception:  # noqa: BLE001 - fall back to the cap if page count is unknown
        total_pages = max_pages
    page_count = min(total_pages, max_pages)

    images: list[bytes] = []
    for page_num in range(1, page_count + 1):
        rendered = convert_from_path(path, dpi=dpi, first_page=page_num, last_page=page_num)
        if not rendered:
            continue
        img = rendered[0].convert("RGB")  # JPEG has no alpha channel
        buf = BytesIO()
        img.save(buf, format="JPEG", quality=quality)
        images.append(buf.getvalue())
        del img, buf, rendered
        gc.collect()
    return images, total_pages


def _ocr_pdf(path: str) -> str:
    """Render each PDF page to a PNG and OCR it with Claude Vision.

    Memory-conscious for the 512MB free tier: render one page at a time at a
    modest DPI, free each rendered image before moving on, and stop after
    _OCR_MAX_PAGES pages.

    Lazy imports keep pdf2image/poppler off the hot path for normal text PDFs
    and avoid an import-time hard dependency. pdf2image needs the system
    `poppler-utils` package (installed in the Dockerfile).
    """
    from pdf2image import convert_from_path, pdfinfo_from_path

    from app.services.claude_service import ocr_image

    try:
        total_pages = pdfinfo_from_path(path)["Pages"]
    except Exception:  # noqa: BLE001 - fall back to the cap if page count is unknown
        total_pages = _OCR_MAX_PAGES
    page_count = min(total_pages, _OCR_MAX_PAGES)

    print(f"[OCR] scanned PDF: {page_count} page(s) to OCR (dpi={_OCR_DPI})", flush=True)
    parts: list[str] = []
    for page_num in range(1, page_count + 1):
        print(f"[OCR] page {page_num}/{page_count}", flush=True)
        # Render just this one page so only a single image is ever in memory.
        images = convert_from_path(
            path, dpi=_OCR_DPI, first_page=page_num, last_page=page_num
        )
        if not images:
            continue
        img = images[0]
        buf = BytesIO()
        img.save(buf, format="PNG")
        page_text = ocr_image(buf.getvalue(), "image/png")
        if page_text.strip():
            parts.append(page_text)
        del img, buf, images
        gc.collect()
    return "\n".join(parts)
